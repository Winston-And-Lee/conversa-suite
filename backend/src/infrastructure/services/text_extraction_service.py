import io
import re
import requests
import fitz  # PyMuPDF for PDF extraction
import docx  # python-docx for DOCX extraction
from fastapi import HTTPException
from typing import Union


class TextExtractionService:
    """Service for extracting text from PDF and DOC files."""

    async def extract_from_url(self, file_url: str, file_type: str) -> str:
        """
        Extract text from a file using its URL.
        
        Args:
            file_url: URL of the file
            file_type: Type of the file (PDF or DOC/DOCX)
            
        Returns:
            str: Extracted text content
        """
        try:
            # Download file from URL
            response = requests.get(file_url)
            if response.status_code != 200:
                raise HTTPException(status_code=500, detail=f"Failed to download file from {file_url}")
            
            file_data = io.BytesIO(response.content)
            
            # Extract text based on file type
            if file_type == "application/pdf":
                return self._extract_from_pdf(file_data)
            elif file_type in ["application/msword", "application/vnd.openxmlformats-officedocument.wordprocessingml.document"]:
                return self._extract_from_docx(file_data)
            else:
                raise HTTPException(status_code=400, detail="Unsupported file type for text extraction")
                
        except Exception as e:
            raise HTTPException(status_code=500, detail=f"Text extraction error: {str(e)}")
    
    def _extract_from_pdf(self, file_data: io.BytesIO) -> str:
        """
        Extract text from a PDF file.
        
        Args:
            file_data: Binary content of the PDF file
            
        Returns:
            str: Extracted text content
        """
        try:
            # Open the PDF file
            doc = fitz.open(stream=file_data, filetype="pdf")
            
            text_content = []
            
            # Extract text from each page
            for page_num in range(len(doc)):
                page = doc.load_page(page_num)
                text_content.append(page.get_text())
            
            # Clean and concatenate text
            full_text = "\n".join(text_content)
            cleaned_text = self._clean_text(full_text)
            
            return cleaned_text
            
        except Exception as e:
            raise Exception(f"PDF extraction error: {str(e)}")
    
    def _extract_from_docx(self, file_data: io.BytesIO) -> str:
        """
        Extract text from a DOCX file.
        
        Args:
            file_data: Binary content of the DOCX file
            
        Returns:
            str: Extracted text content
        """
        try:
            # Open the DOCX file
            doc = docx.Document(file_data)
            
            # Extract text from paragraphs
            paragraphs = [para.text for para in doc.paragraphs]
            
            # Extract text from tables
            tables_text = []
            for table in doc.tables:
                for row in table.rows:
                    row_text = [cell.text for cell in row.cells]
                    tables_text.append(" | ".join(row_text))
            
            # Combine paragraphs and tables
            full_text = "\n".join(paragraphs + tables_text)
            cleaned_text = self._clean_text(full_text)
            
            return cleaned_text
            
        except Exception as e:
            raise Exception(f"DOCX extraction error: {str(e)}")
    
    def _clean_text(self, text: str) -> str:
        """
        Clean extracted text by removing extra whitespace and normalizing.
        
        Args:
            text: Raw extracted text
            
        Returns:
            str: Cleaned text
        """
        # Replace multiple whitespace with a single space
        text = re.sub(r'\s+', ' ', text)
        
        # Remove control characters
        text = re.sub(r'[\x00-\x1F\x7F]', '', text)
        
        # Trim whitespace
        text = text.strip()
        
        return text 